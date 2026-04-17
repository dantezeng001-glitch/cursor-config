#!/usr/bin/env python3
"""
Markdown → DOCX converter.

Usage:
    python md_to_docx.py <input.md> [output.docx]

If output path is omitted, writes to the same directory with .docx extension.
"""

import re
import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Configuration — tweak these to change fonts, sizes, or margins
# ---------------------------------------------------------------------------
FONT_BODY = "Microsoft YaHei"
FONT_CJK = "Microsoft YaHei"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_TABLE = Pt(9.5)
FONT_SIZE_CODE = Pt(9.5)
LINE_SPACING = 1.15
MARGIN = Cm(2.54)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
HEADER_BG = "2B2D42"
TABLE_BORDER_COLOR = "AAAAAA"
QUOTE_BORDER_COLOR = "4A90D9"
CODE_INLINE_COLOR = RGBColor(0x88, 0x00, 0x44)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def setup_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = LINE_SPACING

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_BODY
        hs.font.color.rgb = HEADING_COLOR
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)

    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN


RICH_RE = re.compile(
    r'\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|`(.+?)`'
)


def add_rich_text(paragraph, text: str):
    pos = 0
    for m in RICH_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = FONT_CODE
            run.font.size = FONT_SIZE_CODE
            run.font.color.rgb = CODE_INLINE_COLOR
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def make_table_borders(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def shade_header_row(row):
    for cell in row.cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def parse_table_rows(raw_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in raw_lines:
        line = line.strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        cells = [c.replace("<br>", "\n") for c in cells]
        rows.append(cells)
    return rows


def write_table(doc: Document, raw_lines: list[str]):
    rows = parse_table_rows(raw_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    make_table_borders(table)

    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = row.cells[ci]
            cell.text = ""
            first_para = cell.paragraphs[0]
            first_para.paragraph_format.space_before = Pt(2)
            first_para.paragraph_format.space_after = Pt(2)

            sub_lines = cell_text.split("\n")
            for si, sub in enumerate(sub_lines):
                sub = sub.strip()
                if sub.startswith("* ") or sub.startswith("- "):
                    sub = "• " + sub[2:]
                if not sub:
                    continue
                if si == 0 and not first_para.text and len(first_para.runs) == 0:
                    add_rich_text(first_para, sub)
                else:
                    np = cell.add_paragraph()
                    np.paragraph_format.space_before = Pt(1)
                    np.paragraph_format.space_after = Pt(1)
                    add_rich_text(np, sub)

            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = FONT_SIZE_TABLE

        if ri == 0:
            shade_header_row(row)


def add_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{QUOTE_BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(border)
    add_rich_text(p, text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.italic = True


def add_code_block(doc: Document, code_lines: list[str]):
    for cl in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(cl)
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(input_path: str, output_path: str):
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    i = 0
    table_buf: list[str] = []

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()
        is_table = stripped.startswith("|") and "|" in stripped[1:]

        if is_table:
            table_buf.append(stripped)
            i += 1
            continue
        elif table_buf:
            write_table(doc, table_buf)
            table_buf = []

        if stripped == "" or stripped == "---":
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines):
                cl = lines[i].rstrip("\n")
                if cl.strip().startswith("```"):
                    i += 1
                    break
                code_lines.append(cl)
                i += 1
            add_code_block(doc, code_lines)
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            h = doc.add_heading(level=level)
            add_rich_text(h, heading_match.group(2).strip())
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            add_blockquote(doc, stripped[2:].strip())
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:].strip())
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    if table_buf:
        write_table(doc, table_buf)

    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f"OK  {output_path}  ({size:,} bytes)")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    src = sys.argv[1]
    if not os.path.isfile(src):
        print(f"Error: file not found: {src}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        dst = sys.argv[2]
    else:
        dst = str(Path(src).with_suffix(".docx"))

    convert(src, dst)
